"""Testes de extracao de texto (ETAPA 01B)."""
from pathlib import Path

from licitacao.services.extraction.extract import (
    extrair,
    extrair_docx,
    extrair_pdf,
    extrair_txt,
    extrair_xlsx,
)


def _pdf_textual(path: Path, texto: str, paginas: int = 1):
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path))
    for i in range(paginas):
        c.drawString(72, 720, f"{texto} pagina {i + 1}")
        c.showPage()
    c.save()


def _pdf_vazio(path: Path):
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path))
    c.showPage()
    c.save()


def test_pdf_textual_extrai_pagina():
    arquivo = Path("doc_tmp_pdf1.pdf")
    _pdf_textual(arquivo, "Objeto licitado")
    try:
        r = extrair_pdf(arquivo)
        assert r.erro is None
        assert r.precisa_ocr is False
        assert len(r.trechos) == 1
        assert r.trechos[0].pagina == 1
        assert "Objeto licitado" in r.trechos[0].texto_bruto
    finally:
        arquivo.unlink(missing_ok=True)


def test_pdf_multiplas_paginas_preserva_numero():
    arquivo = Path("doc_tmp_pdf2.pdf")
    _pdf_textual(arquivo, "conteudo", paginas=3)
    try:
        r = extrair_pdf(arquivo)
        assert len(r.trechos) == 3
        assert [t.pagina for t in r.trechos] == [1, 2, 3]
    finally:
        arquivo.unlink(missing_ok=True)


def test_pdf_sem_texto_marca_ocr_required():
    arquivo = Path("doc_tmp_pdf3.pdf")
    _pdf_vazio(arquivo)
    try:
        r = extrair_pdf(arquivo)
        assert r.erro is None
        assert r.precisa_ocr is True
        assert len(r.trechos) == 0
    finally:
        arquivo.unlink(missing_ok=True)


def test_docx_paragrafo_tabela_paragrafo_preserva_ordem(tmp_path: Path):
    from docx import Document

    arquivo = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("Paragrafo um")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "CelulaA"
    t.rows[0].cells[1].text = "CelulaB"
    doc.add_paragraph("Paragrafo dois")
    doc.save(arquivo)

    r = extrair_docx(arquivo)
    assert r.erro is None
    tipos = [t.tipo_localizador for t in r.trechos]
    assert tipos == ["PARAGRAFO", "TABELA", "PARAGRAFO"]
    assert r.trechos[1].tabela == "1"
    assert r.trechos[1].linha_tabela == 1
    assert "CelulaA" in r.trechos[1].texto_bruto


def test_xlsx_coordenadas_reais_ignoram_celula_vazia(tmp_path: Path):
    from openpyxl import Workbook

    arquivo = tmp_path / "plan.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "LOTE 02"
    ws["B7"] = "item"
    ws["D7"] = "preco"
    wb.save(arquivo)

    r = extrair_xlsx(arquivo)
    assert r.erro is None
    assert len(r.trechos) == 1
    t = r.trechos[0]
    assert t.tipo_localizador == "CELULA"
    assert t.planilha == "LOTE 02"
    assert t.celula_inicio == "B7"
    assert t.celula_fim == "D7"


def test_txt_encoding_prioriza_utf8_sig(tmp_path: Path):
    arquivo = tmp_path / "acoes.txt"
    arquivo.write_text("Ação de compra com çedilha e acentuação.", encoding="utf-8-sig")

    r = extrair_txt(arquivo)
    assert r.erro is None
    assert any("Ação" in t.texto_bruto for t in r.trechos)


def test_formato_invalido_retorna_erro(tmp_path: Path):
    arquivo = tmp_path / "arquivo.xyz"
    arquivo.write_bytes(b"xyz")
    r = extrair(arquivo, "xyz")
    assert r.erro is not None


def test_arquivo_corrompido_pdf_retorna_erro(tmp_path: Path):
    arquivo = tmp_path / "corrompido.pdf"
    arquivo.write_bytes(b"este nao e um pdf valido")
    r = extrair_pdf(arquivo)
    assert r.erro is not None
